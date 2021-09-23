#The prototype for the tkinter program

import sys
from tkinter import *
from tkinter.simpledialog import askstring, askinteger, askfloat
from tkinter import messagebox, filedialog
import tkinter.font as font
import docx
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docxcompose.composer import Composer

document = docx.Document()
root = Tk()
root.title("Word Template Creator")

#############################################################################
#Define a function to convert a string into an integer
def convert_to_int(ipstring):
    try:
        value1 = float(ipstring)
        value2 = int(value1)
        if (value1 - value2) == 0.0:
            return value2
        else:
            messagebox.showinfo("Error","Decimal Values Entered")
    except:
        return None

#Define a function to convert a string into a floating point number
def convert_to_float(ipstring):
    try:
        value = float(ipstring)
    except:
        return None
    return value

#Define a function to simply enable a disabled/locked frame
def just_enable(enframe):
    for child in enframe.winfo_children():
        child.configure(state = 'active')

#Define a function to close the overlying widget and enable the main widget based on other values
def enable(widget_to_close, cover = 0, isrecord = 0):
    if (cover == 0) and (isrecord == 0):
        for child in mframe.winfo_children():
            child.configure(state = 'active')
        widget_to_close.destroy()
    else:
        widget_to_close.destroy()

#Define a function to disable a frame        
def disable(frame_to_disable):
    for child in frame_to_disable:
        child.configure(state = 'disable')

#Define a function that displays a message when a user tries to close an important window
def disable_event():
    messagebox.showinfo("Error","Cannot close the window until you've provided the values correctly!")
    pass
    
#Define a funciton to quit the program if the user clicks the x button on the main widget
def quitprogram():
    decision = messagebox.askokcancel("Important!", "Do you really want to exit the application? All progress will not be saved.")
    if decision:
        sys.exit()
    else:
        pass

#Custom message when the user clicks the x button on cover_widget1
def disable_event_cover():
    messagebox.showinfo("Message","Kindly click on the page break button to exit the coverpage building process")
    pass

#Custom fucntion defined for the pagebreak button on the cover_widget2 widget
def cover_pagebreak(document, widget):
    global cover
    global isrecord
    decision = messagebox.askokcancel("Message", "Adding this page break will take you to the main menu. Do you want to continue?")
    if decision:
        cover = 0
        pagebreak(document)
        enable(widget, cover, isrecord)
    else:
        return
        
cover = 0
isrecord = 0
defaults = 0
############################################################################
def coverpage(document):
    global isrecord
    #print("The value of isrecord in coverpage is " + str(isrecord))
    global cover
    cover = 1
    
    #Define the cancel operaiton
    def cancel_operation1():
        global cover
        cover = 0
        enable(cover_widget1, cover, isrecord)
    
    #Get the details of the first widget in the coverpage
    def get_det1():
        namedoc = namedoc_entry.get()
        if (namedoc is None) or (set(namedoc)=={' '}) or (namedoc == ""):
            namedoc = "Enter Title Here"
        #print(namedoc)
        fontname = fontname_entry.get()
        if (fontname is None) or (set(fontname)=={' '}) or (fontname == ""):
            fontname = "Calibri"
        #print(fontname)
        fontsize_str = fontsize_entry.get()
        fontsize = convert_to_float(fontsize_str)
        if fontsize is None:
            fontsize = 38
        #print(fontsize)
        #print(alignment_status.get())
        cover_update1(namedoc, fontname, fontsize, alignment_status.get())
        cover_widget1.destroy()
        cover_widget2_function()
    
    #Update the document with the details obtained from the first widget
    def cover_update1(namedoc, fontname, fontsize, alignment_status):
        headtitle = document.add_heading("",0)
        headrun = headtitle.add_run(namedoc)
        headrun.font.name = fontname
        headrun.font.size = Pt(fontsize)
        if alignment_status == "Left":
            headtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif alignment_status == "Center":
            headtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment_status == "Right":
            headtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment_status == "Justified":
            headtitle.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    #Disable the main widget and define the new widget/window
    disable(mframe.winfo_children())
    #cover_widget1 = Tk()
    cover_widget1 = Toplevel(root)
    cover_widget1.grab_set()
    cover_widget1.geometry("700x250")
    cover_widget1.minsize(700,250)
    cover_widget1.maxsize(700,250)
    cover_widget1.protocol("WM_DELETE_WINDOW", disable_event)
    
    cw1_label = Label(cover_widget1, font = textfont, text = "Create a cover page here! Let's start with these details first").grid(row = 0, column = 0, columnspan = 5)
    
    namedoc_label = Label(cover_widget1, font = labelfont, text = "Enter the title of the coverpage", padx = 10, pady = 10).grid(row = 1, column = 0)
    namedoc_entry = Entry(cover_widget1, borderwidth = 5)
    namedoc_entry.insert(END, "Enter the Title Here")
    namedoc_entry.grid(row = 1, column = 1)
    
    fontname_label = Label(cover_widget1, font = labelfont, text = "Enter the font you want to use", padx = 10, pady = 10).grid(row = 2, column = 0)
    fontname_entry = Entry(cover_widget1, borderwidth = 5)
    fontname_entry.insert(END, "Calibri")
    fontname_entry.grid(row = 2, column = 1)
    
    fontsize_label = Label(cover_widget1, font = labelfont, text = "Enter the font size", padx = 10, pady = 10).grid(row = 3, column = 0)
    fontsize_entry = Entry(cover_widget1, borderwidth = 5)
    fontsize_entry.insert(END, "38")
    fontsize_entry.grid(row = 3, column = 1)
    
    alignment_label = Label(cover_widget1, font = labelfont, text = "What should be the alignment?", padx = 10, pady = 10).grid(row = 4, column = 0)
    alignment_status = StringVar(cover_widget1)
    alignment_status.set("Center")
    alignment_dropdown = OptionMenu(cover_widget1, alignment_status, "Left", "Center", "Right", "Justified").grid(row = 4, column = 1)
    
    submit_button1 = Button(cover_widget1, bg = "OliveDrab1", text = "Submit", padx = 10, pady = 10, command = get_det1).grid(row = 6, column = 0)
    cancel_button1 = Button(cover_widget1, bg = "IndianRed1", text = "Cancel", padx = 10, pady = 10, command = cancel_operation1).grid(row = 6, column = 1)
        
    #Create a widget to perform rest of the actions on the coverpage
    def cover_widget2_function():
        #cover_widget2 = Tk()
        cover_widget2 = Toplevel(root)
        cover_widget2.grab_set()
        cover_widget2.title("Add more to the coverpage")
        cover_widget2.geometry("450x350")
        cover_widget2.minsize(450,350)
        cover_widget2.maxsize(450,350)
        cover_widget2.protocol("WM_DELETE_WINDOW", disable_event_cover)
        
        #cover2_frame  = LabelFrame(cover_widget2, text = "So many options to choose from!", padx = 5, pady = 5, bd = 5)
        #cover2_frame.grid(row = 1, column = 0, padx = 10, pady = 10, columnspan = 5)
        
        cw2_label1 = Label(cover_widget2, font = textfont, text = "Build the rest of the coverpage here!", padx = 10, pady = 20).grid(row = 0, column = 0, columnspan = 5)
        
        add_heading_button = Button(cover_widget2, font = buttonfont, bg = "pale turquoise", pady = 10, text = "Add a Heading", command = lambda: writeheading(document)).grid(row = 1, column = 1)
        
        add_paragraph_button = Button(cover_widget2, font = buttonfont, bg = "cyan", pady = 10, text = "Add a Paragraph", command = lambda: writeparagraph(document)).grid(row = 1, column = 3)
        
        add_table_button = Button(cover_widget2, font = buttonfont, bg = "tan1", padx = 10, pady = 10, text = "Add a Table", command = lambda: addtable(document)).grid(row = 3, column = 1)
        
        add_image_button = Button(cover_widget2, font = buttonfont, bg = "RosyBrown1", padx = 10, pady = 10, text = "Add an Image", command = lambda: addimage(document, cover_widget2)).grid(row = 3, column = 3)
        
        add_pagebreak_button = Button(cover_widget2, font = buttonfont, bg = "steel blue", padx = 10, pady = 10, text = "Add a Page Break", command = lambda: cover_pagebreak(document, cover_widget2)).grid(row = 5, column = 1, columnspan = 3)
        
        fillery1 = Label(cover_widget2, text = "          ", pady = 5).grid(row = 2)
        fillery2 = Label(cover_widget2, text = "          ", pady = 5).grid(row = 4)
    
    cover_widget1.mainloop()
    #print("Widget closed")

def create_margins(document):
    cover = 0
    global isrecord
    #print("The value of isrecord in margins is "+str(isrecord))
    
    #Function to get the details entered in the margins widget
    def get_margin_det():
        defaults = 0
        dtop_margin = convert_to_float(top_margin_entry.get())
        if dtop_margin is None:
            defaults += 1
            dtop_margin = 1
        dbottom_margin = convert_to_float(bottom_margin_entry.get())
        if dbottom_margin is None:
            defaults += 1
            dbottom_margin = 1
        dleft_margin = convert_to_float(left_margin_entry.get())
        if dleft_margin is None:
            defaults += 1
            dleft_margin = 1
        dright_margin = convert_to_float(right_margin_entry.get())
        if dright_margin is None:
            defaults += 1
            dright_margin = 1
        if defaults > 0:
            messagebox.showinfo("Message", "Faulty inputs detected in "+str(defaults)+" places. Using default values in those places")
        enable(margin_widget, cover, isrecord)
        documentmargins(dtop_margin, dbottom_margin, dleft_margin, dright_margin, inchcm.get())

    #Function that actually puts the margins in the document
    def documentmargins(topmargin, bottommargin, leftmargin, rightmargin, inchcm):
        sections_list = document.sections
        if inchcm == "inch":
            sections_list[-1].top_margin = Inches(topmargin)
            sections_list[-1].bottom_margin = Inches(bottommargin)
            sections_list[-1].left_margin = Inches(leftmargin)
            sections_list[-1].right_margin = Inches(rightmargin)
        elif inchcm == "cm":
            sections_list[-1].top_margin = Cm(topmargin)
            sections_list[-1].bottom_margin = Cm(bottommargin)
            sections_list[-1].left_margin = Cm(leftmargin)
            sections_list[-1].right_margin = Cm(rightmargin)
    
    #Define the cancel operaiton
    def cancel_operation():
        enable(margin_widget, cover, isrecord)
        
    #Disable the main widget while margin details are being entered
    disable(mframe.winfo_children())
    proceed = messagebox.askokcancel("Important!", "This may change your default/previous margin selection. Do you want to continue?")
    if proceed:
        #margin_widget = Tk()
        margin_widget = Toplevel(root)
        margin_widget.grab_set()
        margin_widget.title("Set Margins")
        margin_widget.geometry("550x220")
        margin_widget.minsize(550,220)
        margin_widget.maxsize(550,220)
        #Disable the close button as the main widget is still disabled
        margin_widget.protocol("WM_DELETE_WINDOW", disable_event)
        margin_label = Label(margin_widget, font = textfont, text = "Please enter the margins here:", padx = 10, pady = 10).grid(row = 0, column = 0, columnspan = 5)
        #Define the radio buttons
        inchcm = StringVar(margin_widget, "inch")
        Radiobutton(margin_widget, text = "In Inches", variable = inchcm, value = "inch").grid(row = 1, column = 1)
        Radiobutton(margin_widget, text = "In Cm", variable = inchcm, value = "cm").grid(row = 1, column = 2)
        #Design the other entry boxes and labels
        top_margin_label = Label(margin_widget, font = labelfont, text = "Top Margin: ", padx = 5, pady = 5).grid(row = 2, column = 0)
        bottom_margin_label = Label(margin_widget, font = labelfont, text = "Bottom Margin: ", padx = 5, pady = 5).grid(row = 3, column = 0)
        left_margin_label = Label(margin_widget, font = labelfont, text = "Left Margin: ", padx = 5, pady = 5).grid(row = 2, column = 3)
        right_margin_label = Label(margin_widget, font = labelfont, text = "Right Margin: ", padx = 5, pady = 5).grid(row = 3, column = 3)
        top_margin_entry = Entry(margin_widget, borderwidth = 5)
        top_margin_entry.insert(END, "1")
        top_margin_entry.grid(row = 2, column = 1)
        bottom_margin_entry = Entry(margin_widget, borderwidth = 5)
        bottom_margin_entry.insert(END, "1")
        bottom_margin_entry.grid(row = 3, column = 1)
        left_margin_entry = Entry(margin_widget, borderwidth = 5)
        left_margin_entry.insert(END, "1")
        left_margin_entry.grid(row = 2, column = 4)
        right_margin_entry = Entry(margin_widget, borderwidth = 5)
        right_margin_entry.insert(END, "1")
        right_margin_entry.grid(row = 3, column = 4)
        #Submit and Cancel Buttons
        fillery = Label(margin_widget, text = "          ", pady = 5).grid(row = 4)
        margin_submit_button = Button(margin_widget, bg = "OliveDrab1", text = "Submit", padx = 10, pady = 10, command = get_margin_det).grid(row = 5, column = 1)
        margin_cancel_button = Button(margin_widget, bg = "IndianRed1", text = "Cancel", padx = 10, pady = 10, command = cancel_operation).grid(row = 5, column = 3)
    else:
        margin_widget = Tk()
        enable(margin_widget, cover, isrecord)

def addimage(document, widgettoiconify = root): #Use Try Except Blocks with messagebox. Before that try to debug the problem
    widgettoiconify.iconify()
    location = filedialog.askopenfilename(filetypes=[("jpg Images","*.jpg"),("PNG Images","*.png"),("JPEG Images","*.JPEG"), ("JPG Images","*JPG")])
    #print(location)
    if location == "":
        widgettoiconify.deiconify()
        #print("Operation Cancelled")
        return None
    try:
        widgettoiconify.deiconify()
        document.add_picture(location, width = Inches(6))
        document.add_paragraph()
    except:
        messagebox.showinfo("Error!","Cannot add image to the document.")
        return None
        
def writeparagraph(document):
    global isrecord
    global cover
    #print("The value of isrecord in paragraph is " + str(isrecord))
    para_contents = []
    #print("The value of cover in writeparagraph is " + str(cover))
    #Define the function of the cancel button
    def cancel_operation():
        enable(paragraph_widget, cover, isrecord)
        
    #Define the functionality of the next button
    def nextpoint(ipstring):#Update the document here itself. The submit button should be the final command for the paragraph
        global defaults
        if (ipstring == "") or (set(ipstring) == {' '}) or (ipstring == '\n'):
            ipstring = "Write something here "
        para_contents.append(ipstring[:-1])
        textbox.delete("1.0",END)
        paragraph_type = type_of_paragraph.get()
        font_size = convert_to_float(font_size_entry.get())
        if (font_size is None) or (font_size <= 0):
            font_size = 10
            defaults += 1
        font_name = font_name_entry.get()
        if (font_name == None) or (set(font_name) == {' '}) or (font_name == ""):
            font_name = "Georgia"
        bold_format = boldvariable.get()
        italic_format = italicvariable.get()
        underline_format = underlinevariable.get()
        create_paragraph(paragraph_type, font_size, font_name, para_contents[-1], bold_format, italic_format, underline_format)
        #print(para_contents)
    
    #Define the function to get paragraph details
    def get_paragraph_det():
        global cover
        global isrecord
        global defaults
        #print(type_of_paragraph.get())
        #print(boldvariable.get(), italicvariable.get(), underlinevariable.get())
        paragraph_type = type_of_paragraph.get()
        font_size = convert_to_float(font_size_entry.get())
        if (font_size is None) or (font_size <= 0):
            messagebox.showinfo("Error", "Invalid font size entered. Using default value")
            font_size = 10
        if defaults > 0:
            messagebox.showinfo("Error", "Invalid font size entered for some points/paragraphs. Used default value for those cases")
            defaults = 0
        font_name = font_name_entry.get()
        if (font_name == None) or (set(font_name) == {' '}) or (font_name == ""):
            font_name = "Georgia"
        bold_format = boldvariable.get()
        italic_format = italicvariable.get()
        underline_format = underlinevariable.get()
        paragraph_text = textbox.get("1.0", END)
        if (paragraph_text == "") or (set(paragraph_text) == {' '}) or (paragraph_text == "\n"):
            paragraph_text = "Write paragraph here"
        #print(paragraph_text)
        create_paragraph(paragraph_type, font_size, font_name, paragraph_text[:-1], bold_format, italic_format, underline_format)
        enable(paragraph_widget, cover, isrecord)
    
    #Define function that actually creates the paragraph in the document
    def create_paragraph(paragraph_type, font_size, font_name, paragraph_text, bold_format, italic_format, underline_format):
        if paragraph_type == 'normal':
            para = document.add_paragraph()
        elif paragraph_type == 'number':
            para = document.add_paragraph("", style = 'List Number')
        elif paragraph_type == 'bullet':
            para = document.add_paragraph("", style = 'List Bullet')
        else:
            print("There is some error in paragraph type")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pararun = para.add_run(paragraph_text)
        pararun.font.size = Pt(font_size)
        pararun.font.name = font_name
        if bold_format == 1:
            pararun.bold = True
        if italic_format == 1:
            pararun.italic = True
        if underline_format == 1:
            pararun.underline = True
    
    #Disable the main widget while paragraph details are being entered
    disable(mframe.winfo_children())
    #paragraph_widget = Tk()
    paragraph_widget = Toplevel(root)
    paragraph_widget.grab_set()
    paragraph_widget.title("Paragraph Details")
    
    #Mention Dimensions Here and disable the close button
    paragraph_widget.geometry("690x570")
    paragraph_widget.minsize(690,570)
    paragraph_widget.maxsize(690,570)
    paragraph_widget.protocol("WM_DELETE_WINDOW", disable_event)
    
    #Put the elements of the widget here
    paragraph_label = Label(paragraph_widget, font = textfont, text = "Enter the details of the paragraph here!", padx = 10, pady = 10).grid(row = 0, column = 0, columnspan = 3)
        #Define the radio buttons
    type_of_paragraph_label = Label(paragraph_widget, font = labelfont, text = "What is the type of paragraph?", padx = 10, pady = 10).grid(row = 2, column = 0)
    type_of_paragraph = StringVar(paragraph_widget, "normal")
    Radiobutton(paragraph_widget, text = "Normal", variable = type_of_paragraph, value = "normal", tristatevalue = 'x').grid(row = 3, column = 0)
    Radiobutton(paragraph_widget, text = "Number List", variable = type_of_paragraph, value = "number", tristatevalue = 'x').grid(row = 3, column = 1)
    Radiobutton(paragraph_widget, text = "Bullet List", variable = type_of_paragraph, value = "bullet", tristatevalue = 'x').grid(row = 3, column = 2)
        #Define the font size and font name
    font_name_label = Label(paragraph_widget, font = labelfont, text = "Enter the font name:", padx = 10, pady = 10).grid(row = 4, column = 0)
    font_name_entry = Entry(paragraph_widget, borderwidth = 5)
    font_name_entry.insert(END, "Georgia")
    font_name_entry.grid(row = 4, column = 1)
    font_size_label = Label(paragraph_widget, font = labelfont, text = "Enter the font size:", padx = 10, pady = 10).grid(row = 5, column = 0)
    font_size_entry = Entry(paragraph_widget, borderwidth = 5)
    font_size_entry.insert(END, "10")
    font_size_entry.grid(row = 5, column = 1)
        #Define the formatting options checkboxes
    format_label = Label(paragraph_widget, font = labelfont, text = "Choose the formatting options for your paragraph: ", padx = 10, pady = 10).grid(row = 6, column = 0)
    
    boldvariable = IntVar(paragraph_widget, 0)
    italicvariable = IntVar(paragraph_widget, 0)
    underlinevariable = IntVar(paragraph_widget,0)
    
    bold_check1 = Checkbutton(paragraph_widget, text = "Bold", variable = boldvariable, pady = 5)
    bold_check1.grid(row = 7, column = 0)
    
    italic_check1 = Checkbutton(paragraph_widget, text = "Italic", variable = italicvariable, pady = 5)
    italic_check1.grid(row = 7, column = 1)
    
    underline_check1 = Checkbutton(paragraph_widget, text = "Underline", variable = underlinevariable, pady = 5)
    underline_check1.grid(row = 7, column = 2)
    
        #Define the text box for contents
    textbox_label = Label(paragraph_widget, font = labelfont, text = "Enter the contents of the paragraph here: ", padx = 10, pady = 10, anchor = 'n').grid(row = 8, column = 0)
    textbox = Text(paragraph_widget, height = 10, width = 40, pady = 10)
    textbox.grid(row = 8, column = 1, columnspan = 2)
    
        #Define submit, cancel and next buttons
    next_button = Button(paragraph_widget, bg = "PaleTurquoise1", text = "Next Point/Paragraph", pady = 5, command = lambda: nextpoint(textbox.get("1.0", END))).grid(row = 9, column = 2)
    submit_button = Button(paragraph_widget, bg = "OliveDrab1", font = buttonfont, text = "Submit", padx = 10, pady = 10, command = get_paragraph_det).grid(row = 10, column = 0)
    cancel_button = Button(paragraph_widget, bg = "IndianRed1", font = buttonfont, text = "Cancel", padx = 10, pady = 10, command = cancel_operation).grid(row = 10, column = 1)
    filler = Label(paragraph_widget, text = "          ", pady = 15).grid(row = 11)
    
def addtable(document):
    global isrecord
    #print("The value of isrecord in table is " + str(isrecord))
    global cover
    #Define the function of the cancel button
    def cancel_operation():
        enable(table_widget, cover, isrecord)
    
    #Define the function to get table details
    def get_table_det():
        Flag1, Flag2 = True, True
        num_row = convert_to_int(row_entry.get())
        if (num_row is None) or (num_row < 1):
            messagebox.showinfo("Error!","Bad input detected on rows. Please enter a valid positive integer")
            Flag1 = False
        if Flag1 and (num_row > 50):
            messagebox.showinfo("Message","Wow! That's a pretty big table")
        num_col = convert_to_int(col_entry.get())
        if (num_col is None) or (num_col < 1):
            messagebox.showinfo("Error!","Bad input detected on columns. Please enter a valid positive integer")
            Flag2 = False
        if Flag2 and (num_col > 10):
            messagebox.showinfo("Problem","Too many columns given. Entering 10 columns only")
            num_col = 10
        if (Flag1 is False) or (Flag2 is False):
            table_widget.destroy()
            addtable(document)
        if Flag1 and Flag2:
            enable(table_widget, cover, isrecord)
            createtable(num_row+1, num_col)
        
    #Define function that actually creates the table in the document
    def createtable(nrows, ncols):
        table = document.add_table(nrows, ncols)
        table.style = 'Table Grid'
        for c in range(ncols):
            table.cell(0, c).text = "Col "+ str(c + 1) + " Heading"
        document.add_paragraph()
    
    #Disable the main widget while table details are being entered
    disable(mframe.winfo_children())
    #table_widget = Tk()
    table_widget = Toplevel(root)
    table_widget.grab_set()
    table_widget.title("Set the Table")
    #Mention Dimensions Here and disable the close button
    table_widget.geometry("620x300")
    table_widget.minsize(620,200)
    table_widget.maxsize(620,200)
    table_widget.protocol("WM_DELETE_WINDOW", disable_event)
    #Put the elements of the widget here
    table_label = Label(table_widget, font = textfont, text = "Enter number rows and columns of the table here:", padx = 10, pady = 10).grid(row = 0, column = 0, columnspan = 4)
    row_label = Label(table_widget, font = labelfont, text = "Please enter the number of rows: ", padx = 10, pady = 5).grid(row = 1, column = 0)
    col_label = Label(table_widget, font = labelfont, text = "Please enter the number of columns: ", padx = 8, pady = 5).grid(row = 2, column = 0)
    row_entry = Entry(table_widget, borderwidth = 5)
    row_entry.grid(row = 1, column = 1)
    col_entry = Entry(table_widget, borderwidth = 5)
    col_entry.grid(row = 2, column = 1)
    fillery = Label(table_widget, text = "          ", pady = 5).grid(row = 3)
    submit_button = Button(table_widget, bg = "OliveDrab1", text = "Submit", padx = 10, pady = 10, command = get_table_det).grid(row = 4, column = 0)
    cancel_button = Button(table_widget, bg = "IndianRed1", text = "Cancel", padx = 10, pady = 10, command = cancel_operation).grid(row = 4, column = 1)
    
def writeheading(document):
    global isrecord
    #print("The value of isrecord in heading is " + str(isrecord))
    #Define the function of the cancel button
    def cancel_operation():
        enable(heading_widget, cover, isrecord)
    
    #Define the function to get heading widget details
    def get_heading_det():
        global cover
        global isrecord
        heading_name = heading_name_entry.get()
        if (heading_name == "") or (set(heading_name) == {' '}) or (heading_name == None):
            heading_name = "Write heading here"
        level = levelvar.get()
        bold_format = boldvariable.get()
        italic_format = italicvariable.get()
        underline_format = underlinevariable.get()
        enable(heading_widget, cover, isrecord)
        create_heading(heading_name, level, bold_format, italic_format, underline_format)
    
    #Define function that actually creates the heading in the document
    def create_heading(heading_name, level, bold_format, italic_format, underline_format):
        heading = document.add_heading("", level)
        headrun = heading.add_run(heading_name)
        if bold_format == 1:
            headrun.font.bold = True
        if italic_format == 1:
            headrun.font.italic = True
        if underline_format == 1:
            headrun.font.underline = True
    
    #Enter the details of the widget here
        #Disable the main widget while heading details are being entered
    disable(mframe.winfo_children())
    #heading_widget = Tk()
    heading_widget = Toplevel(root)
    heading_widget.grab_set()
    heading_widget.title("Heading Details")
        
        #Mention Dimensions Here and disable the close button
    heading_widget.geometry("400x250")
    heading_widget.minsize(400,250)
    heading_widget.maxsize(400,250)
    heading_widget.protocol("WM_DELETE_WINDOW", disable_event)
        
        #Put the elements of the widget here
    heading_label = Label(heading_widget, font = textfont, text = "Details of the Heading:", padx = 10, pady = 10).grid(row = 0, column = 0, columnspan = 5)
    heading_name_label = Label(heading_widget, font = labelfont, text = "Enter the Heading here: ", padx = 10, pady = 10).grid(row = 1, column = 0)
    heading_name_entry = Entry(heading_widget, borderwidth = 5)
    heading_name_entry.insert(END, "Enter Heading Here")
    heading_name_entry.grid(row = 1, column = 1)
        
        #Add Dropdown for Level
    levelvar = IntVar(heading_widget)
    levelvar.set(1)
    level_label = Label(heading_widget, font = labelfont, text = "Level:", padx = 10, pady = 10).grid(row = 2, column = 0)
    level_drop = OptionMenu(heading_widget, levelvar, 0, 1, 2, 3, 4, 5, 6, 7, 8, 9).grid(row = 2, column = 1)
        
        #Add checkboxes for Bold, Italic, and Underlines
    format_label = Label(heading_widget, font = labelfont, text = "Choose the formatting that you want for the heading: ", padx = 10, pady = 10).grid(row = 3, column = 0, columnspan = 3)
    
    boldvariable = IntVar(heading_widget, 0)
    italicvariable = IntVar(heading_widget, 0)
    underlinevariable = IntVar(heading_widget,0)
    
    bold_check1 = Checkbutton(heading_widget, text = "Bold", variable = boldvariable, pady = 5)
    bold_check1.grid(row = 4, column = 0)
    
    italic_check1 = Checkbutton(heading_widget, text = "Italic", variable = italicvariable, pady = 5)
    italic_check1.grid(row = 4, column = 1)
    
    underline_check1 = Checkbutton(heading_widget, text = "Underline", variable = underlinevariable, pady = 5)
    underline_check1.grid(row = 4, column = 2)
    
        #Submit and Cancel Buttons
    submit_button = Button(heading_widget, bg = "OliveDrab1", text = "Submit", padx = 5, pady = 5, command = get_heading_det).grid(row = 6, column = 0)
    cancel_button = Button(heading_widget, bg = "IndianRed1", text = "Cancel", padx = 5, pady = 5, command = cancel_operation).grid(row = 6, column = 1)
    
def pagebreak(document):    #Explore dialog boxes for either section breaks or page breaks
    document.add_page_break()
    messagebox.showinfo("Message!", "Added a page break at the end of the document")

def record(document):
    global cover
    global isrecord
    isrecord = 1
    messagebox.showinfo("Note", "The instructions that you give now will be recorded!")
    #Create a new document
    doc1 = docx.Document()
    
    #A special step before choosing margins for a recorded section
    def record_create_margins():
        decision = messagebox.askokcancel("Important!", "Please note that adding a margin here creates a new section and may cause a page break if previous margins don't align. Do you want to continue?")
        if decision:
            document.add_section()
            create_margins(document)
        else:
            return
            
    #Functionality after stopping the recording
    def stoprecord():
        global isrecord
        finalstep = messagebox.askokcancel("Important!","Do you want to stop recording instructions?")
        if finalstep:
            try:
                num_repeat = askinteger("Times to Repeat", "How many times do you want to repeat the recorded instructions?")
                if (num_repeat < 0) or (num_repeat == None):
                    messagebox.showinfo("Alert!","Bad input detected. Please provide a positive integer value")
                    stoprecord()
                composer = Composer(document)
                for i in range(num_repeat):
                    composer.append(doc1)
                isrecord = 0
            #print("The value of isrecord after record is " + str(isrecord))
                enable(record_window, cover, isrecord)
            except:
                print("Passed exception")
                stoprecord()
        else:
            return
    
    #Create a widget that houses all the functions that can be recorded
        #Disable the main widget while heading details are being entered
    disable(mframe.winfo_children())
    #record_window = Tk()
    record_window = Toplevel(root)
    record_window.grab_set()
    root.iconify()
    record_window.title("Let's make a record!")
    record_window.geometry("550x550")
    record_window.minsize(530,550)
    record_window.maxsize(530,550)
    record_window.protocol("WM_DELETE_WINDOW", disable_event)
    
    record_label = Label(record_window, font = textfont, text = 'Choose any functionality', padx = 10, pady = 10).grid(row = 0, column = 1, columnspan = 3)
    
    #record_frame  = LabelFrame(record_window, text = "So many options to choose from!", padx = 5, pady = 5, bd = 5)
    #record_frame.grid(row = 1, column = 0, padx = 10, pady = 10, columnspan = 5)
    
    record_coverpage_button = Button(record_window, bg = "azure", font = buttonfont, text = "Add a \nCoverpage", padx = 10, pady = 10, command = lambda: coverpage(doc1)).grid(row = 1, column = 1)
    record_margins_button = Button(record_window, bg = "gray69", font = buttonfont, text = "Add Margins", padx = 30, pady = 20, command = record_create_margins).grid(row = 1, column = 3)
    fillerx1 = Label(record_window, padx = 10, pady = 10, text = "          ").grid(row = 2)
    record_image_button = Button(record_window, bg = "RosyBrown1", font = buttonfont, text = "Add an \nImage", padx = 30, pady = 10, command = lambda: addimage(doc1, record_window)).grid(row = 3, column = 1)
    record_paragraph_button = Button(record_window, bg = "cyan", font = buttonfont, text = "Add a \nParagraph", padx = 40, pady = 10, command = lambda: writeparagraph(doc1)).grid(row = 3, column = 3)
    fillerx2 = Label(record_window, padx = 10, pady = 10, text = "          ").grid(row = 4)
    record_table_button = Button(record_window, bg = "tan1", font = buttonfont, text = "Add a \nTable", padx = 35, pady = 10, command = lambda: addtable(doc1)).grid(row = 5, column = 1)
    record_heading_button = Button(record_window, bg = "pale turquoise", font = buttonfont, text = "Add a \nHeading", padx = 52, pady = 10, command = lambda: writeheading(doc1)).grid(row = 5, column = 3)
    fillerx3 = Label(record_window, padx = 10, pady = 10, text = "          ").grid(row = 6)
    fillery1 = Label(record_window, padx = 10, pady = 10, text = "          ").grid(column = 0)
    fillery2 = Label(record_window, padx = 10, pady = 10, text = "          ").grid(column = 2)
    fillery3 = Label(record_window, padx = 10, pady = 10, text = "          ").grid(column = 4)
    record_break_button = Button(record_window, bg = "steel blue", font = buttonfont, text = "Add a \nPage Break", padx = 10, pady = 10, command = lambda: pagebreak(doc1)).grid(row = 7, column = 1)
    stop_record_button = Button(record_window, bg = "orchid1", font = buttonfont, text = "Stop Recording \nInstructions", padx = 16, pady = 10, command = stoprecord).grid(row = 7, column = 3)
    
    record_window.mainloop()

#Saving the document - #Change the isalnum functionality. Try implementing asksavefileas
def savedoc(document):
    decision = messagebox.askokcancel("Important!", "This will save your changes and close the program. Do you want to continue?")
    if decision:
        name = askstring("Save As","What is the name of the file?                    ")
        if name == None:
            messagebox.showinfo("Message", "Operation cancelled")
            return None
            name = "New Word File"
        finalname = ""
        for i in name:
            if i in ['\\','/',':','*','?','"','<','>','|']:
                messagebox.showinfo("Name Error!", "The name cannot contain " + str(i) + "charcter. Omitting the character")
            else:
                finalname += i
        if finalname == "":
            finalname = "New Word Document"
        savename = finalname + ".docx"
        try:
            document.save("../../Created Files/"+ savename)
        except:
            document.save("../../"+savename)
        messagebox.showinfo("Program Complete!", "Your file has been saved as "+ finalname +".docx")
        root.destroy()
    else:
        pass

textfont = font.Font(family = "Times New Roman", size = 20, weight = 'bold', slant = 'italic') 
welcome_message = Label(root, text = "Welcome to the Word Template Creator!", font = textfont).grid(row = 0, column = 1, columnspan = 5)
about_message = Label(root, text = "You can do a lot of things here! Click on any of the buttons to start building your document!", pady = 15, font = textfont).grid(row = 1, column = 1, columnspan = 5)

#Defining the main frame
mframe  = LabelFrame(root, text = "So many options to choose from!", padx = 5, pady = 5, bd = 5)
mframe.grid(row = 2, column = 0, padx = 10, pady = 10, columnspan = 5)

#Defining the fonts
buttonfont = font.Font(family = "Arial", size = 16, weight = "bold", slant = "italic")
labelfont = font.Font(family = "Garamond", size = 10, weight = 'bold')

#Defining the buttons
coverpage_button = Button(mframe, text = 'Add Coverpage', font = buttonfont, bg = "azure", padx = 50, pady = 30, command = lambda: coverpage(document)).grid(row = 2, column = 1)
margins_button = Button(mframe, text = 'Add Margins', font = buttonfont, bg = "gray69", padx = 50, pady = 30, command = lambda: create_margins(document)).grid(row = 2, column = 3)
images_button = Button(mframe, text = 'Add Image', font = buttonfont, bg = "RosyBrown1", padx = 66, pady = 30, command = lambda: addimage(document)).grid(row = 2, column = 5)
paragraph_button = Button(mframe, text = 'Add Paragraph', font = buttonfont, bg = "cyan", padx = 51, pady = 30, command = lambda: writeparagraph(document)).grid(row = 4, column = 1)
table_button = Button(mframe, text = 'Add a Table', font = buttonfont, bg = "tan1", padx = 53, pady = 30, command = lambda: addtable(document)).grid(row = 4, column = 3)
heading_button = Button(mframe, text = 'Add a Heading', font = buttonfont, bg = "pale turquoise",padx = 50, pady = 30, command = lambda: writeheading(document)).grid(row = 4, column = 5)
page_break_button = Button(mframe, text = 'Add Page Break', font = buttonfont, bg="steel blue", padx = 48, pady = 30, command = lambda: pagebreak(document)).grid(row = 6, column = 1)
record_button = Button(mframe, text = 'Record Instructions', font = buttonfont, bg = "orchid1", padx = 20, pady = 30, command = lambda: record(document)).grid(row = 6, column = 3)
save_button = Button(mframe, text = 'Save Document', font = buttonfont, bg = "medium spring green", padx = 46, pady = 30, command = lambda:savedoc(document)).grid(row = 6, column = 5)
filler1 = Label(mframe, text = "          ").grid(column = 2)
filler2 = Label(mframe, text = "          ").grid(column = 4)
filler3 = Label(mframe, text = "          ").grid(column = 6)
filler4 = Label(mframe, text = "          ", pady = 15).grid(row = 3)
filler5 = Label(mframe, text = "          ", pady = 15).grid(row = 5)
root.protocol("WM_DELETE_WINDOW", quitprogram)
root.mainloop()