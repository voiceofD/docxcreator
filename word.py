#Program to create automated Word Templates/Documents

#Import the dependencies
import docx
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docxcompose.composer import Composer

#Define the document
document = docx.Document()

#Define the cover page
def coverpage(document):
    alignment_flag = False
    namedoc = input('Enter the Title of the coverpage: ')
    headtitle = document.add_heading("",0)
    headrun = headtitle.add_run(namedoc)
    try:
        headrun.font.name = input('Enter the font you want to use (Default: Calibri): ') or 'Calibri'
        headtitlefontsize = int(input('Enter the font size (Default: 38): ')) or 38
    except:
        print("Invalid Input detected. Using default value")
        headtitlefontsize = 38
        headrun.font.name = 'Calibri'
    headrun.font.size = Pt(headtitlefontsize)
    while True:
        if alignment_flag == False:
            alignment_decision = input('Enter the alignment (Default: Center): Left(1), Center(2), Right(3), Justified(4): ') or "2"
            if alignment_decision not in ['1','2','3','4']:
                print("Please enter a valid option")
                continue
            else:
                alignment_flag = True
        format = input('Mention B,I, or U to Bold, Italicize or Underline (Press Enter to do nothing): ') or 'cool'
        if len(format) > 10:
            print("Please provide a valid option (3 or fewer letters)")
            continue
        for j in set(format.lower()):
            if j == 'b':
                headrun.font.bold = True
            if j == 'i':
                headrun.font.italic = True
            if j == 'u':
                headrun.font.underline = True
        break
    if alignment_decision == '1':
        headtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif alignment_decision == '2':
        headtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment_decision == '3':
        headtitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        headtitle.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    while True:
        choice = input('Do you want to add a heading (1), paragraph (2), table (3), image (4), or page break (5): ')
        if choice not in ['1','2','3','4','5']:
            print("Please enter a valid choice")
            continue
        if choice == "1":
            writeheading(document)
        elif choice == "2":
            writeparagraph(document)
        elif choice == '3':
            addtable(document)
        elif choice == '4':
            addimage(document)
        else:
            pagebreak(document)
            break

#Create margins
def create_margins(document):
    while True:
        decision = input('This may change your default/previous margin selection. Do you want to continue? (y/n): ') or 'y'
        if decision.lower() in ['y','yes','yup','yeah','yea','ok','affirmative','right','okay','yep','cool']:
            break
        elif decision.lower() in ['n','no','nope','nuh','na','naw','negative']:
            return None
        else:
            print("I did not understand your command. Please try again")
    sections_list = document.sections
    f1,f2 = False, False
    while True:
        if f1 == False:
            inchcm = input("Margins in Inches(1) or Cm (2)?: ")
            if inchcm not in ["1","2"]:
                print("Please provide valid input")
                continue
            else:
                f1 = True
        if f2 == False:
            try:
                topmargin = float(input("Enter the top margin: "))
                bottommargin = float(input("Enter the bottom margin: "))
                leftmargin = float(input("Enter the left margin: "))
                rightmargin = float(input("Enter the right margin: "))
                if (topmargin < 0) or (bottommargin < 0) or (leftmargin < 0)or (rightmargin < 0):
                    print('Please enter valid values')
                    continue
                else:
                    f2 = True
            except:
                print("Please enter valid numeric values")
                continue
        break
    if inchcm == "1":
        sections_list[-1].top_margin = Inches(topmargin)
        sections_list[-1].bottom_margin = Inches(bottommargin)
        sections_list[-1].left_margin = Inches(leftmargin)
        sections_list[-1].right_margin = Inches(rightmargin)
    else:
        sections_list[-1].top_margin = Cm(topmargin)
        sections_list[-1].bottom_margin = Cm(bottommargin)
        sections_list[-1].left_margin = Cm(leftmargin)
        sections_list[-1].right_margin = Cm(rightmargin)

#Define adding an image
def addimage(document):
    while True:
        try:
            location = input("Path to the image: ") or "test.jpg"
            document.add_picture(location, Inches(7), Inches(4))
            document.add_paragraph()
            break
        except:
            print('Path not found')
            continue

#Write a paragraph in the document
def writeparagraph(document):
    while True:
        type_of_paragraph = input('What type of paragraph is this? (Normal(1), Number List(2), Bullet List(3)): ')
        if type_of_paragraph not in ['1','2','3']:
            print("Please enter a valid option")
            continue
        break
    content = input("What is the content of the paragraph?: ") or "Write something here"
    if type_of_paragraph == '1':
        para = document.add_paragraph()
    elif type_of_paragraph == '2':
        para = document.add_paragraph("", style = 'List Number')
    else:
        para = document.add_paragraph("", style = 'List Bullet')
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pararun = para.add_run(content)
    try:
        pararunfontsize = int(input('Enter the font size (Default: 10): ')) or 10
        pararun.font.name = input('Enter the font name (Default: Georgia): ') or "Georgia"
    except:
        print("Invalid input detected. Using default value")
        pararunfontsize = 10
        pararun.font.name = 'Georgia'
    pararun.font.size = Pt(pararunfontsize)
    while True:
        format = input('Mention B,I, or U to Bold, Italicize or Underline (Press Enter to do nothing): ') or '1'
        if len(format) > 10:
            print("Please provide a valid option (3 or fewer letters)")
            continue
        for j in set(format.lower()):
            if j == 'b':
                pararun.bold = True
            if j == 'i':
                pararun.italic = True
            if j == 'u':
                pararun.underline = True
        break
    
#Define a table
def addtable(document):
    while True:
        try:
            r = int(input("How many rows?: "))
            if r > 50:
                print("That's a pretty big table!")
            c = int(input("How many columns?: "))
            if c > 10:
                print("Too many columns. Try again")
                continue
            if (r <= 0) or (c <= 0):
                print('Please enter valid values')
                continue
        except:
            print('Please enter valid values')
            continue
        break
    table = document.add_table(r,c)
    table.style = 'Table Grid'
    for col in range(c):
        table.cell(0, col).text = "Col "+ str(col + 1) +" Heading"
    document.add_paragraph()

#Write a heading in the document
def writeheading(document):
    headingname = input('Enter name of the heading: ') or "Write Heading Here"
    while True:
        try:
            level = int(input('What should be the level of the heading?: '))
            if (level < 0) or (level > 9):
                print('Enter a valid integer from 0 to 9')
                continue
        except:
            print('Please enter valid integers from 0 to 9')
            continue
        break
    heading = document.add_heading("", level)
    headrun = heading.add_run(headingname)
    #Ask the user if they want to bold, italicize, or underline the heading
    while True:
        format = input('Mention B,I, or U to Bold, Italicize or Underline (Press Enter to do nothing): ') or 'cool'
        if len(format) > 10:
            print("Please provide a valid option (3 or fewer letters)")
            continue
        for j in set(format.lower()):
            if j == 'b':
                headrun.font.bold = True
            if j == 'i':
                headrun.font.italic = True
            if j == 'u':
                headrun.font.underline = True
        break

#Define a page break
def pagebreak(document):
    document.add_page_break()

#Record instructions and repeat it n times
def record(document):
    doc1 = docx.Document()
    print("The instructions that you type now will be recorded")
    while True:
        print("The function codes are as follows:\n1) Coverpage\n2) Margins\n3) Add Image\n4) Write Paragraph\n5) Add Table\n6) Write Heading\n7) Page Break\n8) Stop recording instructions")
        try:
            x = int(input("Enter the function code: "))
        except:
            print('Please enter a valid input')
            continue
        if x == 1:
            coverpage(doc1)
        elif x == 2:
            print("Please note that adding a margin here creates a new section and may cause a page break if previous margins don't align")
            document.add_section()
            create_margins(document)
        elif x == 3:
            addimage(doc1)
        elif x == 4:
            writeparagraph(doc1)
        elif x == 5:
            addtable(doc1)
        elif x == 6:
            writeheading(doc1)
        elif x == 7:
            pagebreak(doc1)
        elif x == 8:
            toggle = False
            break
        else:
            print("Function not yet defined")
    print("The instructions have been recorded")
    while True:
        try:
            n = int(input("How many times do you want to repeat the recorded instructions?: "))
            if n < 1:
                print("Please provide a valid input")
                continue
        except:
            print("Please provide a valid input")
        break
    composer = Composer(document)
    for i in range(n):
        composer.append(doc1)

#Save the document
def savedoc(document):
    name = input('What should the file be saved as?: ') or "New File"
    finalname = ""
    for i in name:
        if i in ['\\','/',':','*','?','"','<','>','|']:
            continue
        else:
            finalname += i
    if finalname == "":
        finalname = "New File"
    document.save(name + ".docx")
    
#Run the program
while True:
    print("The function codes are as follows:\n1) Coverpage\n2) Margins\n3) Add Image\n4) Write Paragraph\n5) Add Table\n6) Write Heading\n7) Page Break\n8) Record Instructions\n9) Save")
    try:
        x = int(input("Enter the function code: "))
        if (x <= 0):
            print("Please provide a valid value")
    except:
        print('Please enter a valid input')
        continue
    if x == 1:
        coverpage(document)
    elif x == 2:
        create_margins(document)
    elif x == 3:
        addimage(document)
    elif x == 4:
        writeparagraph(document)
    elif x == 5:
        addtable(document)
    elif x == 6:
        writeheading(document)
    elif x == 7:
        pagebreak(document)
    elif x == 8:
        record(document)
    elif x == 9:
        savedoc(document)
        break
    else:
        print("Function not yet defined")
        
''' The above snippet is dependent on the following additional packages: python-docx; docxcompose;
The following 'bugs' have been observed
--> The margins bug: The margins defined after recorded section will affect the recorded section: Can be rectified by adding a section break functionality (Or record instruction only 1 time to define margin)
The scope for further development are as follows:
--> Add features for section breaks, page borders, font colors, headers/footers, enter table values, edit existing documents'''
